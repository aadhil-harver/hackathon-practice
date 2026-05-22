"""CV file → text extraction.

Used by the Streamlit file-upload widget. Supports the three formats recruiters
actually receive in practice — PDF, DOCX, and plain text. Scanned (image-only)
PDFs are not handled here; they'd need an OCR pass (tesseract / pytesseract)
which is out of scope for the hackathon demo.

The extractor returns plain UTF-8 text, lightly cleaned of common artifacts
(NUL bytes, redundant whitespace) but otherwise unmodified. The parser agent
downstream is responsible for understanding structure.
"""

from __future__ import annotations

import io
import re
from typing import Protocol


class UploadedLike(Protocol):
    """Duck-type for Streamlit's ``UploadedFile`` so this module stays UI-agnostic."""

    name: str

    def read(self) -> bytes: ...
    def seek(self, offset: int) -> int: ...


_WS_RE = re.compile(r"[ \t]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")


class CVExtractionError(RuntimeError):
    """Raised when an upload cannot be parsed into usable text."""


def _clean(text: str) -> str:
    """Strip NUL bytes, collapse runs of whitespace, normalise blank-line gaps."""
    text = text.replace("\x00", "")
    text = _WS_RE.sub(" ", text)
    text = _BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader  # local import keeps Streamlit boot light

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise CVExtractionError(f"Could not open PDF: {exc}") from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            # Skip pages that fail (encrypted/malformed) but keep going so a
            # mostly-readable CV still yields useful text.
            continue

    text = _clean("\n\n".join(p for p in parts if p))
    if not text:
        raise CVExtractionError(
            "No text extracted. This looks like a scanned/image-only PDF — "
            "the demo doesn't OCR. Re-export the CV with a text layer."
        )
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise CVExtractionError(f"Could not open DOCX: {exc}") from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Tables (lots of CVs put dates / skills in tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = _clean("\n".join(paragraphs))
    if not text:
        raise CVExtractionError("DOCX opened but contained no readable text.")
    return text


def _extract_txt(data: bytes) -> str:
    text = _clean(data.decode("utf-8", errors="replace"))
    if not text:
        raise CVExtractionError("Text file was empty.")
    return text


def extract_text(upload: UploadedLike) -> str:
    """Dispatch to the right extractor based on file extension.

    The caller (Streamlit) should already have validated the upload's type via
    ``st.file_uploader(type=[...])`` — this is a second guard.
    """
    name = (upload.name or "").lower()
    upload.seek(0)
    data = upload.read()

    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt"):
        return _extract_txt(data)
    raise CVExtractionError(
        f"Unsupported file type: {name!r}. Use .pdf, .docx, or .txt."
    )
